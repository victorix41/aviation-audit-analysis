"""Generate management-ready Word reports for aviation audit analytics."""

from datetime import date
from io import BytesIO

import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.ui.data_service import (
    ExecutiveOverviewResult,
    generate_executive_overview,
)

REPORT_TITLE = "Aviation MRO Audit Findings Management Report"


def _set_document_defaults(
    document: DocumentObject,
) -> None:
    """Apply basic report formatting."""

    section = document.sections[0]

    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10)


def _add_title_page(
    document: DocumentObject,
    *,
    source_file_name: str,
    as_of_date: date,
) -> None:
    """Add the report title and reporting details."""

    title = document.add_heading(
        REPORT_TITLE,
        level=0,
    )

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run("Management Decision-Support Report")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)

    document.add_paragraph()

    details = document.add_table(
        rows=3,
        cols=2,
    )

    details.style = "Table Grid"

    details.cell(0, 0).text = "Source file"
    details.cell(0, 1).text = source_file_name

    details.cell(1, 0).text = "As-of date"
    details.cell(1, 1).text = as_of_date.strftime("%d %B %Y")

    details.cell(2, 0).text = "Report basis"
    details.cell(2, 1).text = "Cleaned and validated aviation audit register"

    document.add_paragraph()

    notice = document.add_paragraph()

    notice_run = notice.add_run("Important qualification: ")
    notice_run.bold = True

    notice.add_run(
        "Past-due classifications are based on the recorded response "
        "due date. The dataset does not currently contain a finding "
        "status or closure date, so a past-due date does not prove that "
        "a finding remains open."
    )

    document.add_page_break()


def _add_executive_kpis(
    document: DocumentObject,
    overview: ExecutiveOverviewResult,
) -> None:
    """Add executive KPI tables."""

    summary = overview.audit_summary

    document.add_heading(
        "1. Executive Overview",
        level=1,
    )

    document.add_heading(
        "Finding Profile",
        level=2,
    )

    finding_table = document.add_table(
        rows=2,
        cols=4,
    )

    finding_table.style = "Table Grid"

    finding_headers = [
        "Total Findings",
        "Major",
        "Minor",
        "Observations",
    ]

    finding_values = [
        summary.total_findings,
        summary.major_count,
        summary.minor_count,
        summary.observation_count,
    ]

    for index, header in enumerate(finding_headers):
        finding_table.cell(0, index).text = header
        finding_table.cell(1, index).text = str(finding_values[index])

    document.add_paragraph()

    document.add_heading(
        "Response Due-Date Position",
        level=2,
    )

    due_date_table = document.add_table(
        rows=2,
        cols=4,
    )

    due_date_table.style = "Table Grid"

    due_date_headers = [
        "Past Due",
        "Due Within 30 Days",
        "Future Due",
        "Missing Due Date",
    ]

    due_date_values = [
        summary.past_due_response_count,
        summary.due_within_30_days_count,
        summary.future_due_count,
        summary.missing_due_date_count,
    ]

    for index, header in enumerate(due_date_headers):
        due_date_table.cell(0, index).text = header
        due_date_table.cell(1, index).text = str(due_date_values[index])


def _add_leading_categories(
    document: DocumentObject,
    overview: ExecutiveOverviewResult,
) -> None:
    """Add leading human-factor and root-cause results."""

    document.add_heading(
        "2. Leading Contributing Categories",
        level=1,
    )

    category_table = document.add_table(
        rows=3,
        cols=4,
    )

    category_table.style = "Table Grid"

    headers = [
        "Category",
        "Leading Result",
        "Frequency",
        "Percentage",
    ]

    for index, header in enumerate(headers):
        category_table.cell(0, index).text = header

    human_factor = overview.human_factor_analysis

    category_table.cell(1, 0).text = "Human Factor"
    category_table.cell(1, 1).text = human_factor.top_factor or "Not available"
    category_table.cell(1, 2).text = str(human_factor.top_factor_frequency)
    category_table.cell(1, 3).text = f"{human_factor.top_factor_percentage:.2f}%"

    root_cause = overview.root_cause_analysis

    category_table.cell(2, 0).text = "Root Cause"
    category_table.cell(2, 1).text = root_cause.top_root_cause or "Not available"
    category_table.cell(2, 2).text = str(root_cause.top_root_cause_frequency)
    category_table.cell(2, 3).text = f"{root_cause.top_root_cause_percentage:.2f}%"


def _add_severity_table(
    document: DocumentObject,
    overview: ExecutiveOverviewResult,
) -> None:
    """Add severity Pareto results."""

    document.add_heading(
        "3. Severity Distribution",
        level=1,
    )

    pareto_table = overview.severity_analysis.pareto.table

    if pareto_table.empty:
        document.add_paragraph("No severity information is available.")
        return

    table = document.add_table(
        rows=1,
        cols=4,
    )

    table.style = "Table Grid"

    headers = [
        "Severity",
        "Frequency",
        "Percentage",
        "Cumulative Percentage",
    ]

    for index, header in enumerate(headers):
        table.cell(0, index).text = header

    for _, row in pareto_table.iterrows():
        cells = table.add_row().cells

        cells[0].text = str(row["category"])
        cells[1].text = str(int(row["frequency"]))
        cells[2].text = f"{float(row['percentage']):.2f}%"
        cells[3].text = f"{float(row['cumulative_percentage']):.2f}%"


def _add_monthly_workload(
    document: DocumentObject,
    overview: ExecutiveOverviewResult,
) -> None:
    """Add monthly response workload results."""

    document.add_heading(
        "4. Monthly Response Workload",
        level=1,
    )

    monthly_trend = overview.severity_analysis.monthly_trend

    if monthly_trend.empty:
        document.add_paragraph(
            "No valid response due dates are available for monthly workload analysis."
        )
        return

    displayed_columns = [
        column
        for column in [
            "period",
            "Observation",
            "Minor",
            "Major",
            "Unspecified",
            "total",
        ]
        if column in monthly_trend.columns
    ]

    table = document.add_table(
        rows=1,
        cols=len(displayed_columns),
    )

    table.style = "Table Grid"

    for index, column in enumerate(displayed_columns):
        table.cell(0, index).text = column.replace(
            "_",
            " ",
        ).title()

    for _, row in monthly_trend.iterrows():
        cells = table.add_row().cells

        for index, column in enumerate(displayed_columns):
            value = row[column]

            if column == "period":
                cells[index].text = str(value)
            else:
                cells[index].text = str(int(value))

    document.add_paragraph(
        "Trend qualification: this table groups findings by "
        "response due date. It represents response workload, not "
        "the date each finding originally occurred."
    )


def _add_management_insights(
    document: DocumentObject,
    overview: ExecutiveOverviewResult,
) -> None:
    """Add deterministic observations and recommendations."""

    insights = overview.executive_insights

    document.add_heading(
        "5. Management Observations",
        level=1,
    )

    if insights.has_observations:
        for observation in insights.observations:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(observation)
    else:
        document.add_paragraph("No management observations were generated.")

    document.add_heading(
        "6. Management Considerations",
        level=1,
    )

    if insights.has_recommendations:
        for recommendation in insights.recommendations:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(recommendation)
    else:
        document.add_paragraph(
            "No additional management considerations were generated."
        )

    document.add_paragraph()

    qualification = document.add_paragraph()

    qualification_run = qualification.add_run("Method note: ")
    qualification_run.bold = True

    qualification.add_run(
        "The observations and considerations in this report are "
        "generated using deterministic rules from the displayed "
        "analytics. They are not AI-generated conclusions."
    )


def generate_management_word_report(
    dataframe: pd.DataFrame,
    *,
    source_file_name: str,
    as_of_date: date,
) -> bytes:
    """
    Generate a management-ready Word report.

    Args:
        dataframe:
            Cleaned aviation audit dataset.

        source_file_name:
            Name of the uploaded audit register.

        as_of_date:
            Date used for response-due-date classification.

    Returns:
        Word document content as bytes.
    """

    overview = generate_executive_overview(
        dataframe,
        as_of_date=as_of_date,
    )

    document = Document()

    _set_document_defaults(document)

    _add_title_page(
        document,
        source_file_name=source_file_name,
        as_of_date=as_of_date,
    )

    _add_executive_kpis(
        document,
        overview,
    )

    document.add_paragraph()

    _add_leading_categories(
        document,
        overview,
    )

    document.add_paragraph()

    _add_severity_table(
        document,
        overview,
    )

    document.add_paragraph()

    _add_monthly_workload(
        document,
        overview,
    )

    document.add_paragraph()

    _add_management_insights(
        document,
        overview,
    )

    output = BytesIO()
    document.save(output)

    return output.getvalue()
