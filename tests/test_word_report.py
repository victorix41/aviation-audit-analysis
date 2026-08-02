"""Tests for the Word management report generator."""

from datetime import date
from io import BytesIO

import pandas as pd
from docx import Document

from src.reports.word_report import (
    REPORT_TITLE,
    generate_management_word_report,
)


def create_cleaned_dataframe() -> pd.DataFrame:
    """Create cleaned audit data for report testing."""

    return pd.DataFrame(
        {
            "severity_level": [
                "Major",
                "Minor",
                "Observation",
            ],
            "response_due_date": pd.to_datetime(
                [
                    "2026-01-10",
                    "2026-03-10",
                    "2026-04-10",
                ]
            ),
            "human_factor": [
                "Knowledge gap",
                "Time pressure",
                "Knowledge gap",
            ],
            "root_cause": [
                "Procedure weakness",
                "Training gap",
                "Procedure weakness",
            ],
        }
    )


def test_generate_management_word_report() -> None:
    report_bytes = generate_management_word_report(
        create_cleaned_dataframe(),
        source_file_name="audit-register.xlsx",
        as_of_date=date(2026, 2, 1),
    )

    assert report_bytes.startswith(b"PK")
    assert len(report_bytes) > 1000


def test_word_report_contains_expected_headings() -> None:
    report_bytes = generate_management_word_report(
        create_cleaned_dataframe(),
        source_file_name="audit-register.xlsx",
        as_of_date=date(2026, 2, 1),
    )

    document = Document(BytesIO(report_bytes))

    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert REPORT_TITLE in document_text
    assert "Executive Overview" in document_text
    assert "Management Observations" in document_text
    assert "Management Considerations" in document_text


def test_word_report_contains_source_file() -> None:
    report_bytes = generate_management_word_report(
        create_cleaned_dataframe(),
        source_file_name="audit-register.xlsx",
        as_of_date=date(2026, 2, 1),
    )

    document = Document(BytesIO(report_bytes))

    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "audit-register.xlsx" in table_text
    assert "01 February 2026" in table_text
